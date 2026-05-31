"""Word 文档创建工具：根据文本内容生成 .docx 文件。"""

from __future__ import annotations

import os
from typing import Any, Dict, List

from tools.base import Tool, ToolParameter, UserRefusedError
from tools.confirm import confirm_in_cli, get_allow_all_windows_cmd


class CreateDocxTool(Tool):
    """创建 Word 文档（.docx）。"""

    def __init__(self) -> None:
        super().__init__(
            name="create_docx",
            description=(
                "创建 Word 文档（.docx）。"
                "支持标题和正文内容，按换行拆分为多个段落。"
            ),
            expandable=False,
        )

    def get_parameters(self) -> List[ToolParameter]:
        return [
            ToolParameter(
                name="path",
                type="string",
                description="输出 docx 文件路径，例如 docs/report.docx。",
                required=True,
            ),
            ToolParameter(
                name="content",
                type="string",
                description="正文内容；按换行拆分为段落。",
                required=True,
            ),
            ToolParameter(
                name="title",
                type="string",
                description="可选标题，若提供会作为文档一级标题。",
                required=False,
            ),
        ]

    def run(self, parameters: Dict[str, Any]) -> str:
        try:
            from docx import Document  # type: ignore
        except ImportError:
            return "缺少依赖 `python-docx`，请先执行: pip install python-docx"

        path = str(parameters.get("path", "")).strip()
        content = str(parameters.get("content", ""))
        title = str(parameters.get("title", "")).strip()

        if not path:
            return "path 不能为空。"
        if not content.strip() and not title:
            return "content 与 title 不能同时为空。"

        if not path.lower().endswith(".docx"):
            path = path + ".docx"

        abs_path = os.path.abspath(path)
        parent = os.path.dirname(abs_path)

        if not get_allow_all_windows_cmd():
            approved = confirm_in_cli(
                "工具: create_docx\n"
                f"目标文件: {abs_path}\n"
                f"正文行数: {len(content.splitlines())}"
            )
            if not approved:
                raise UserRefusedError("create_docx", f"用户拒绝创建文档: {abs_path}")

        try:
            if parent:
                os.makedirs(parent, exist_ok=True)

            doc = Document()
            if title:
                doc.add_heading(title, level=1)

            for line in content.splitlines():
                doc.add_paragraph(line)

            if not content.splitlines() and title:
                doc.add_paragraph("")

            doc.save(abs_path)
        except OSError as e:
            return f"写入失败: {e}"
        except Exception as e:
            return f"创建 Word 文档失败: {e}"

        para_count = len(content.splitlines()) + (1 if title else 0)
        return f"已创建 Word 文档: {abs_path}（约 {para_count} 段）"
